from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "CodexDocs" / "EONAPP_Investor_Presentation_Deutsch_2026-05-29_FINAL.docx"
SHOT_DIR = ROOT / "docs" / "qa" / "launch-signoff" / "screenshots" / "german-presentation-2026-05-29"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ("sz", "val", "color", "space"):
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def format_cell(cell, bold=False, size=10.5, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.LEFT):
    for p in cell.paragraphs:
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        for r in p.runs:
            r.font.name = "Calibri"
            r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
            r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
            r.font.size = Pt(size)
            r.font.bold = bold
            r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    run = paragraph.add_run("Seite ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def set_run_font(run, name="Calibri", size=11, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text="", style=None, align=None, before=0, after=6, size=11, color="1B1F2A", bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color, italic=italic)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.2 * level)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(8 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4 if level == 1 else 3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.add_run(text)
    return p


def add_callout(doc, title, body, fill="F3F6FB"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.4)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    set_cell_border(cell, top={"sz": 12, "val": "single", "color": "C8D4E3"}, bottom={"sz": 12, "val": "single", "color": "C8D4E3"}, left={"sz": 12, "val": "single", "color": "C8D4E3"}, right={"sz": 12, "val": "single", "color": "C8D4E3"})
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + "\n")
    set_run_font(r1, size=11.5, bold=True, color="0B2545")
    r2 = p.add_run(body)
    set_run_font(r2, size=10.5, color="1B1F2A")
    return table


def add_screen_section(doc, title, image_path, notes):
    add_heading(doc, title, 2)
    add_paragraph(
        doc,
        "Dieser Screenshot stammt aus der live veröffentlichten Version. Bitte prüfen Sie Layout, Lesbarkeit, Navigation, deutsche Beschriftungen und ob die Oberfläche ohne technische Erklärungen verständlich bleibt.",
        after=5,
        size=10.5,
        color="334155",
    )
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(6.35))
        p.paragraph_format.space_after = Pt(4)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    r = cap.add_run(f"Abbildung: {title} ({image_path.name})")
    set_run_font(r, size=9.5, italic=True, color="64748B")
    for note in notes:
        add_bullet(doc, note)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("EONAPP.ch – Investor- und Testpräsentation · ")
    set_run_font(r, size=8.5, color="64748B")
    add_page_number(p)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, attr, Inches(1.0))
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    for name, size, color in [("Title", 24, "0B2545"), ("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "1F4D78"), ("Heading 3", 12, "1F4D78")]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    add_footer(section)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("EONAPP.ch")
    set_run_font(r, size=12, bold=True, color="F97316")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(10)
    r = p2.add_run("Investor- und Testerpräsentation")
    set_run_font(r, size=24, bold=True, color="0B2545")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(12)
    r = p3.add_run("Deutsch · Live-Version · Screenshot-basierter Rundgang · Schritt-für-Schritt-Testplan")
    set_run_font(r, size=12.5, color="475569")
    add_callout(
        doc,
        "Kurzfassung",
        "EONAPP ist ein KI-Business-Cockpit für nicht technische Nutzer. Die App ist live, mehrsprachig, lokal-first und so gebaut, dass der Nutzer in natürlicher Sprache arbeiten kann. Für die Grundbenutzung sind keine API-Schlüssel nötig; die Sprache der Oberfläche funktioniert ohne Zusatzkonfiguration.",
        fill="EEF4FF",
    )

    add_paragraph(doc, "", after=0)
    info = doc.add_table(rows=2, cols=4)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    widths = [Inches(1.35), Inches(2.1), Inches(1.35), Inches(1.55)]
    for row in info.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w
    headers = ["Dokument", "Zielgruppe", "Version", "Datum"]
    values = ["EONAPP Investor/Test Präsentation", "Tester, Partner, Investoren", "Final", "2026-05-29"]
    for i, h in enumerate(headers):
        cell = info.cell(0, i)
        cell.text = h
        set_cell_shading(cell, "D9E8F5")
        format_cell(cell, bold=True, size=10.5, color="0B2545", align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, v in enumerate(values):
        cell = info.cell(1, i)
        cell.text = v
        set_cell_shading(cell, "FFFFFF")
        format_cell(cell, size=10.2, color="111827", align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in info.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"sz": 8, "val": "single", "color": "CBD5E1"}, bottom={"sz": 8, "val": "single", "color": "CBD5E1"}, left={"sz": 8, "val": "single", "color": "CBD5E1"}, right={"sz": 8, "val": "single", "color": "CBD5E1"})
    doc.add_paragraph()

    add_heading(doc, "1. Worum es bei EONAPP geht", 1)
    add_paragraph(doc, "EONAPP.ch ist kein traditionelles Einzweck-Tool. Die Plattform ist als KI-Business-Cockpit gedacht, also als zentraler Arbeitsraum für Chat, Browser, Vault, Inhalte, Marktanalyse, Projekte und geführte Automationen. Die Oberfläche ist bewusst so gebaut, dass man sich nicht durch viele Technik-Menüs kämpfen muss.", after=5)
    add_bullet(doc, "Der Nutzer beschreibt das Ziel in natürlicher Sprache.")
    add_bullet(doc, "EONBOT hilft dabei, Aufgaben zu verstehen, zu planen und zu begleiten.")
    add_bullet(doc, "Die App ist lokal-first, mehrsprachig und ohne Pflicht zur zentralen Identität nutzbar.")
    add_bullet(doc, "Für viele Grundfunktionen braucht man keinen API-Schlüssel.")

    add_heading(doc, "2. Was der Tester konkret prüfen soll", 1)
    add_bullet(doc, "Sind die Seiten sofort verständlich, auch ohne technisches Vorwissen?")
    add_bullet(doc, "Ist die deutsche Oberfläche wirklich lesbar und konsistent?")
    add_bullet(doc, "Funktioniert der Sprachwechsel auf allen Seiten?")
    add_bullet(doc, "Sind wichtige Buttons sichtbar und klar beschriftet?")
    add_bullet(doc, "Treffen Ladezeit und visuelle Qualität einen professionellen Eindruck?")
    add_bullet(doc, "Treten beim Öffnen der Seiten Fehler, leere Bereiche oder kryptische Meldungen auf?")

    add_heading(doc, "3. Testmatrix der wichtigsten Seiten", 1)
    matrix = doc.add_table(rows=1, cols=4)
    matrix.alignment = WD_TABLE_ALIGNMENT.CENTER
    matrix.autofit = False
    matrix.columns[0].width = Inches(1.4)
    matrix.columns[1].width = Inches(1.5)
    matrix.columns[2].width = Inches(2.4)
    matrix.columns[3].width = Inches(1.5)
    headers = ["Seite", "Wofür sie da ist", "Worauf achten", "Erwartung"]
    for i, h in enumerate(headers):
        cell = matrix.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "D9E8F5")
        format_cell(cell, bold=True, size=10, color="0B2545", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_border(cell, top={"sz": 8, "val": "single", "color": "CBD5E1"}, bottom={"sz": 8, "val": "single", "color": "CBD5E1"}, left={"sz": 8, "val": "single", "color": "CBD5E1"}, right={"sz": 8, "val": "single", "color": "CBD5E1"})
    rows = [
        ("Startseite", "Einstieg in die Plattform", "Klarer Einstieg, sichtbare Hauptaktionen", "Verständlich in 5 Sekunden"),
        ("KI-Cockpit", "Hauptarbeitsoberfläche", "Navigation, Aufgabenfluss, EONBOT", "Ein Ort für die Arbeit"),
        ("KI-Chat", "Dialog mit der KI", "Sprache, Hilfe, Vorschläge", "Chat wirkt intelligent und ruhig"),
        ("Vault", "Identität, Schlüssel, Proofs", "Übersicht, lokale Speicherung, Rückkehr", "Vertrauenswürdig und übersichtlich"),
        ("Trade", "Marktanalyse", "Charts, Übersicht, keine Fehler", "Professioneller Produktbereich"),
        ("Market", "Angebote, Kaufen, Verdienen", "Farben, Buttons, Lesbarkeit", "Verkaufsfläche mit Stil"),
        ("Creator Studio", "Inhalte, Video, Musik", "Werkzeuge, Klarheit, Export", "Kreativ und leistungsfähig"),
        ("Reward Access", "Zugang und Belohnungen", "Keine weiße Stub-Seite", "Richtige Produktseite"),
        ("Tools", "Agenten-Werkzeuge", "Nützliche, klare Tool-Kacheln", "Hilfreich statt leer"),
        ("Realm / Profile", "Öffentliche Präsenz", "Klarer Profilpfad, Rollen", "Verständlich und glaubwürdig"),
        ("Onboarding", "Erster Einstieg", "Sprache, Hilfe, erste Schritte", "Einfach für Nicht-Techniker"),
        ("Free AI Power", "Lokale KI/Provider-Erkennung", "Ollama, Modelle, Status", "Einfacher Einstieg ohne Technikstress"),
    ]
    for row_data in rows:
        cells = matrix.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = value
            set_cell_shading(cells[i], "FFFFFF")
            format_cell(cells[i], size=9.7, color="111827", align=WD_ALIGN_PARAGRAPH.LEFT if i in (0, 1, 2) else WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_border(cells[i], top={"sz": 6, "val": "single", "color": "D1D5DB"}, bottom={"sz": 6, "val": "single", "color": "D1D5DB"}, left={"sz": 6, "val": "single", "color": "D1D5DB"}, right={"sz": 6, "val": "single", "color": "D1D5DB"})
    doc.add_paragraph()

    add_heading(doc, "4. Schritt-für-Schritt-Testplan für einen nicht technischen Tester", 1)
    steps = [
        "Die Startseite in Chrome öffnen und prüfen, ob die Hauptaussage sofort verständlich ist.",
        "Die Sprache auf Deutsch stellen und beobachten, ob die komplette Oberfläche deutsch bleibt.",
        "Zu KI-Chat, KI-Cockpit, Vault, Market, Trade und Creator Studio wechseln und jeweils die Überschrift lesen.",
        "In jedem Bereich 1–2 Buttons anklicken und schauen, ob die Seite logisch reagiert.",
        "Die Sprache testweise auf Französisch, Japanisch oder Arabisch umstellen und prüfen, ob die Seite weiterhin lesbar bleibt.",
        "Den Language-Switch auf jeder Seite suchen und prüfen, ob man damit wieder zurück auf Deutsch kommt.",
        "Die Ladezeit einschätzen: Wirkt die Seite schnell und direkt oder zäh und blockiert?",
        "Die Konsole öffnen (F12) und prüfen, ob sichtbare Fehler erscheinen.",
        "Falls Ollama vorhanden ist: Die lokale KI-Erkennung in Free AI Power prüfen.",
        "Am Ende kurzes Feedback geben: Was war klar? Was war verwirrend? Was fehlte? Was wirkt besonders stark?",
    ]
    for idx, step in enumerate(steps, 1):
        add_number(doc, step)

    doc.add_page_break()
    add_heading(doc, "5. Live-Screenshots der wichtigsten Seiten", 1)
    add_paragraph(doc, "Alle folgenden Abbildungen stammen aus der live veröffentlichten Version von EONAPP.ch und wurden im deutschen Sprachmodus aufgenommen.", after=8, color="334155")

    screenshots = [
        ("Startseite", SHOT_DIR / "01-home-index.png", ["Das zentrale Hero-Layout zeigt das KI-Business-Cockpit.", "Die Hauptnavigation ist oben sichtbar.", "Deutsch ist auf den ersten Blick lesbar."]),
        ("KI-Cockpit", SHOT_DIR / "02-ai-cockpit.png", ["Hier erkennt man den Einstieg in den Cockpit-Arbeitsraum.", "Die Seite wirkt wie ein operatives Hauptpanel.", "Die Sprach- und EONBOT-Hilfen bleiben sichtbar."]),
        ("KI-Chat", SHOT_DIR / "03-chat.png", ["Chat ist als Hauptassistent positioniert.", "Der Nutzer kann im Alltag in natürlicher Sprache arbeiten.", "Die Seite soll sich wie ein vertrauter KI-Chat anfühlen."]),
        ("Workbench", SHOT_DIR / "04-workbench.png", ["Die Workbench bündelt Aufgaben, Missionen und Arbeitsmodi.", "Der Fokus liegt auf Ausführung statt nur Ausgabe.", "Gut für Nutzer, die strukturiert arbeiten möchten."]),
        ("Creator Studio", SHOT_DIR / "05-creator-studio.png", ["Hier liegen Inhalte, Video und kreative Workflows.", "Die Seite soll für Creator und Business-Teams verständlich sein.", "Die Oberfläche darf leistungsfähig wirken, aber nicht überladen."]),
        ("Vault", SHOT_DIR / "06-vault.png", ["Der Vault ist die lokale Identitäts- und Beweisfläche.", "Er zeigt Profile, Schlüssel, Badges und lokale Zustände.", "Nicht-technische Nutzer sehen hier ihre Daten übersichtlich gesammelt."]),
        ("Market", SHOT_DIR / "07-market.png", ["Der Marktbereich soll wie eine echte Produkt- und Verkaufsfläche wirken.", "Farben, Buttons und Struktur sollen professionell sein.", "Wichtig: kein weißer Stub-Eindruck."]),
        ("Trade", SHOT_DIR / "08-trade.png", ["Hier wird Marktanalyse und Trading-Information gezeigt.", "Die Seite darf fachlich wirken, aber nicht chaotisch.", "Die Farben und Kontraste müssen gut lesbar sein."]),
        ("Creator Market / Marketplace", SHOT_DIR / "09-marketplace.png", ["Diese Seite verbindet Kreative, Angebote und Marktplatz-Flows.", "Wichtig ist eine klare visuelle Hierarchie.", "Die Seite darf nicht leer oder technisch aussehen."]),
        ("Reward Access", SHOT_DIR / "10-reward-access.png", ["Zugang, Belohnungen und Freischaltungen werden hier gezeigt.", "Die Seite muss wie ein echter Produktbereich aussehen.", "Weißer Leerraum oder Stub-Eindruck wären hier falsch."]),
        ("Tools", SHOT_DIR / "11-tools.png", ["Der Tool-Hub bündelt Agenten-Werkzeuge im Hintergrund.", "Er soll nützlich und aufgeräumt wirken.", "Kein Dummy-Dashboard: hier sollen echte Fähigkeiten sichtbar werden."]),
        ("Onboarding", SHOT_DIR / "12-onboarding.png", ["Das Onboarding hilft neuen Nutzern beim Einstieg.", "Die Sprache muss hier besonders klar und freundlich sein.", "Hier merkt man, ob die App wirklich anfängerfreundlich ist."]),
        ("Leaderboard", SHOT_DIR / "13-leaderboard.png", ["Die Belohnungs- und Referral-Logik wird hier erklärt.", "Die Seite soll motivierend und nicht überladen sein.", "Milestones und Pool Points müssen klar erkennbar sein."]),
        ("Realm", SHOT_DIR / "14-realm.png", ["Realm zeigt Identität, öffentliche Präsenz und Profilräume.", "Die Seite muss vertrauenswürdig und klar wirken.", "Auch hier sollte Deutsch sauber lesbar sein."]),
        ("About", SHOT_DIR / "15-about.png", ["Die About-Seite erklärt die Produktidee und Architektur.", "Sie ist wichtig für Verständnis und Vertrauen.", "Der Text sollte auch für Partner lesbar sein."]),
        ("Privacy", SHOT_DIR / "16-privacy.png", ["Die Privacy-Seite erklärt lokale Datenhaltung und Schutz.", "Sie soll seriös, ruhig und verständlich sein.", "Wichtig für Vertrauen bei neuen Nutzern."]),
        ("Subscription", SHOT_DIR / "17-subscription.png", ["Hier sieht man die Pläne, Limits und die Monetarisierung.", "Die Preise und Begriffe müssen verständlich sein.", "Die Seite sollte als klare Business-Seite wirken."]),
        ("Free AI Power", SHOT_DIR / "18-free-ai-power.png", ["Die Seite zeigt lokale KI und Provider-Erkennung.", "Ideal für Nutzer ohne Technikkenntnisse.", "Ollama und lokale Modelle sollen auffindbar sein."]),
        ("Code Maker", SHOT_DIR / "19-code-maker.png", ["Der Website-/App-Builder zeigt die technische Produktstärke.", "Wichtig: trotzdem verständlich bleiben.", "Die Seite soll nach echter Produktivität aussehen, nicht nach Demo."]),
    ]
    for title, img, notes in screenshots:
        add_screen_section(doc, title, img, notes)

    add_heading(doc, "6. Wie man KI-Schlüssel hinzufügt", 1)
    add_paragraph(doc, "Für die Grundnutzung ist kein API-Schlüssel nötig. Wenn Sie aber einen Anbieter verwenden möchten, gehen Sie in den KI-Bereich oder in den Vault/API-Bereich und tragen dort Ihren Schlüssel ein. Die App sollte den Anbieter dann automatisch verwenden, ohne dass Sie jeden technischen Schritt kennen müssen.", after=4)
    add_bullet(doc, "Öffnen Sie KI-Chat oder KI-Cockpit.")
    add_bullet(doc, "Wählen Sie den gewünschten Anbieter oder Modus aus.")
    add_bullet(doc, "Fügen Sie den Schlüssel nur hinzu, wenn Sie einen bereits besitzen.")
    add_bullet(doc, "Die deutsche Oberfläche bleibt auch ohne Schlüssel nutzbar.")

    add_heading(doc, "7. Lokale KI mit Ollama", 1)
    add_paragraph(doc, "Wenn auf dem Gerät Ollama installiert ist, sollte EONAPP verfügbare lokale Modelle erkennen. Das ist für Nicht-Techniker besonders wichtig, weil die App dadurch sofort zeigen kann, welche Modelle auf dem eigenen Rechner sinnvoll nutzbar sind.", after=4)
    add_bullet(doc, "Lokale KI ist ideal, wenn Sie ohne externe API-Kosten starten möchten.")
    add_bullet(doc, "Die App sollte passende kleine Modelle erkennen und anzeigen.")
    add_bullet(doc, "Wenn kein lokales Modell vorhanden ist, soll die Oberfläche trotzdem sauber bleiben.")
    add_bullet(doc, "Die Nutzerführung darf nicht mit Fachbegriffen überladen werden.")

    add_heading(doc, "8. Worauf der Tester am Ende ein Feedback geben soll", 1)
    feedback = doc.add_table(rows=1, cols=2)
    feedback.alignment = WD_TABLE_ALIGNMENT.CENTER
    feedback.autofit = False
    feedback.columns[0].width = Inches(2.0)
    feedback.columns[1].width = Inches(4.4)
    for i, h in enumerate(["Thema", "Frage an den Tester"]):
        cell = feedback.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "D9E8F5")
        format_cell(cell, bold=True, size=10, color="0B2545", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_border(cell, top={"sz": 8, "val": "single", "color": "CBD5E1"}, bottom={"sz": 8, "val": "single", "color": "CBD5E1"}, left={"sz": 8, "val": "single", "color": "CBD5E1"}, right={"sz": 8, "val": "single", "color": "CBD5E1"})
    feedback_rows = [
        ("Sprache", "War Deutsch überall verständlich und vollständig?"),
        ("Navigation", "Haben die Menüpunkte klar erklärt, wohin sie führen?"),
        ("Design", "Wirkte eine Seite leer, kaputt oder zu technisch?"),
        ("Tempo", "Fühlte sich die Seite schnell und direkt an?"),
        ("Fehler", "Gab es sichtbare Fehlermeldungen oder verwirrende Texte?"),
        ("Zielgruppe", "Würde ein nicht technischer Nutzer die App verstehen?"),
    ]
    for a, b in feedback_rows:
        row = feedback.add_row().cells
        row[0].text = a
        row[1].text = b
        set_cell_shading(row[0], "FFFFFF")
        set_cell_shading(row[1], "FFFFFF")
        format_cell(row[0], size=9.8, color="111827", align=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(row[1], size=9.8, color="111827", align=WD_ALIGN_PARAGRAPH.LEFT)
        for cell in row:
            set_cell_border(cell, top={"sz": 6, "val": "single", "color": "D1D5DB"}, bottom={"sz": 6, "val": "single", "color": "D1D5DB"}, left={"sz": 6, "val": "single", "color": "D1D5DB"}, right={"sz": 6, "val": "single", "color": "D1D5DB"})

    add_paragraph(doc, "", after=0)
    add_callout(
        doc,
        "Abschluss",
        "Bitte prüfen Sie die Seite wie ein echter Anwender: ohne Fachwissen, ohne die Entwicklerbrille, aber mit dem Blick auf Verständlichkeit, Qualität und Vertrauen. Genau dafür ist diese Präsentation gedacht.",
        fill="F4F7FB",
    )

    doc.save(str(OUT))
    print(f"Saved to {OUT}")


if __name__ == '__main__':
    main()
